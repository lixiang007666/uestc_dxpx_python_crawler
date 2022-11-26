import requests
from bs4 import BeautifulSoup
from pypinyin import lazy_pinyin
import docx
import csv
import pandas as pd
judge = {'A': '正确', 'B': '错误'}

# cookie用你们自己的，我这不提供
headers = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                         'Chrome/95.0.4638.69 Safari/537.36',
           'cookie':''
           }
# 主页url地址
father_url = 'https://dxpx.uestc.edu.cn'
#
first_url = 'https://dxpx.uestc.edu.cn/jjfz/exam_center/record?l_id=254'
# 综合测试地址
final_url = 'https://dxpx.uestc.edu.cn/jjfz/exam_center/end_show?rid=343358'


def one_page(url):
    # 使用lxml解析页面代码
    r = requests.get(url=url, headers=headers)
    soup = BeautifulSoup(r.text, "html.parser")

    # 两次定位，先找到整个信息区域，需先手动打开网页源码查看
    info_list = soup.find_all(attrs={'class': 'error_sub'})
    for i in range(len(info_list)):
        if i < 10:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            exam_list = info_list[i].find_all(attrs={'class': 'exam_result2'})
            result_list = info_list[i].find_all(attrs={"class": 'sub_color'})

            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【单选题】', '').replace(str(i + 1) + '、',
                                                                                                       '')
            exam = exam_list[0].text.replace(' ', '').replace('\n', '')
            result = result_list[0].text.replace('正确答案：', '')

            dan_list.append(title + '   ' + exam + "【" + result + '】')

        elif i >= 10 and i < 15:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            exam_list = info_list[i].find_all(attrs={'class': 'exam_result_box2'})
            result_list = info_list[i].find_all(attrs={"class": 'sub_color'})

            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【多选题】', '').replace(str(i + 1) + '、',
                                                                                                       '')
            exam = exam_list[0].text.replace(' ', '').replace('\n', '')
            result = result_list[0].text.replace('正确答案：', '')

            duo_list.append(title + '   ' + exam + "【" + result + '】')
        else:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            result_list = info_list[i].find_all(attrs={"class": 'sub_color'})
            # print(title_list[0].text.replace(' ', '').replace('\n', '').replace('【判断题】','').replace(str(i+1)+'、',''))
            # print(judge[str(result_list[0].text.replace('正确答案：', '').replace(' ', '').replace('\n', ''))])
            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【判断题】', '').replace(str(i + 1) + '、',
                                                                                                       '')
            result = judge[str(result_list[0].text.replace('正确答案：', '').replace(' ', '').replace('\n', ''))]

            panduan_list.append(title + '   ' + "【" + result + '】')


def final(url):
    r = requests.get(url=url, headers=headers)
    soup = BeautifulSoup(r.text, "html.parser")
    # 两次定位，先找到整个信息区域，需先手动打开网页源码查看
    info_list = soup.find_all(attrs={'class': 'error_sub'})
    for i in range(100):
        if i < 30:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            exam_list = info_list[i].find_all(attrs={'class': 'exam_result2'})
            result_list = info_list[i].find_all(attrs={"class": 'sub_color'})

            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【单选题】', '').replace(str(i + 1) + '、',
                                                                                                       '')
            exam = exam_list[0].text.replace(' ', '').replace('\n', '')
            result = result_list[0].text.replace('正确答案：', '')

            dan_list.append(title + '   ' + exam + "【" + result + '】')

        elif i >= 30 and i < 60:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            exam_list = info_list[i].find_all(attrs={'class': 'exam_result_box2'})
            result_list = info_list[i].find_all(attrs={"class": 'sub_color'})

            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【多选题】', '').replace(str(i + 1) + '、',
                                                                                                       '')
            exam = exam_list[0].text.replace(' ', '').replace('\n', '')
            result = result_list[0].text.replace('正确答案：', '')

            duo_list.append(title + '   ' + exam + "【" + result + '】')
        elif i >= 60 and i < 80:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            result_list = info_list[i].find_all(attrs={"class": 'sub_color'})
            # print(title_list[0].text.replace(' ', '').replace('\n', '').replace('【判断题】','').replace(str(i+1)+'、',''))
            # print(judge[str(result_list[0].text.replace('正确答案：', '').replace(' ', '').replace('\n', ''))])
            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【判断题】', '').replace(str(i + 1) + '、',
                                                                                                       '')
            result = judge[str(result_list[0].text.replace('正确答案：', '').replace(' ', '').replace('\n', ''))]
            panduan_list.append(title + '   ' + "【" + result + '】')
        else:
            title_list = info_list[i].find_all(attrs={'class': 'sub_title'})
            result_list = info_list[i].find_all(attrs={'class':'sub_cont'})
            result=result_list[0].text.replace('\n', '')
            title = title_list[0].text.replace(' ', '').replace('\n', '').replace('【填空题】', '').replace(str(i + 1) + '、','').replace('（）','【'+str(result)+'】').replace('（\xa0\xa0）','【'+str(result)+'】').replace('《》','【'+str(result)+'】')
            tiankong_list.append(title)


if __name__ == '__main__':

    url_list = []
    dan_list = []
    duo_list = []
    panduan_list = []
    tiankong_list=[]
    r = requests.get(url=first_url, headers=headers)
    if r.status_code != 200:
        print("访问失败")
    # 使用lxml解析页面代码
    soup = BeautifulSoup(r.text, "html.parser")
    url_list = soup.find_all(attrs={'class': 'error_span_a1'})
    # 爬取10次章节测试题库
    for url_last in url_list:
        one_url = father_url + str(url_last['href'])
        one_page(one_url)
    # 爬取综合测试题库
    final(final_url)
    dan_sort = sorted(dan_list, key=lambda i: lazy_pinyin(i[0]))
    duo_sort = sorted(duo_list, key=lambda i: lazy_pinyin(i[0]))
    panduan_sort = sorted(panduan_list, key=lambda i: lazy_pinyin(i[0]))
    tiankong_sort = sorted(tiankong_list, key=lambda i: lazy_pinyin(i[0]))
    file = docx.Document()
    j=1
    for str1 in dan_sort:
        file.add_paragraph(str(j)+'、'+str1)
        j=j+1
    file.add_paragraph('-----------------------多选-------------------------')
    j = 1
    for str1 in duo_sort:
        file.add_paragraph(str(j)+'、'+str1)
        j = j + 1
    file.add_paragraph('-----------------------判断------------------------')
    j = 1
    for str1 in panduan_sort:
        file.add_paragraph(str(j)+'、'+str1)
        j = j + 1
    j = 1
    for str1 in tiankong_sort:
        file.add_paragraph(str(j)+'、'+str1)
        j = j + 1
    file.save(r"D:\题库.docx")
    #
    # f=open(r"D:\题库.csv",'w',encoding='utf-8')
    # csv_writer = csv.writer(f)
    # for str in dan_sort:
    #     csv_writer.writerow(str)
    # file.add_paragraph('-----------------------多选-------------------------')
    # for str in duo_sort:
    #     csv_writer.writerow(str)
    # file.add_paragraph('-----------------------判断------------------------')
    # for str in panduan_sort:
    #     csv_writer.writerow(str)
    # file.add_paragraph('-----------------------填空------------------------')
    # for str in tiankong_sort:
    #     csv_writer.writerow(str)


